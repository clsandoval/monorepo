"""
Generate Word (.docx) boundary survey reports.

Report structure:
  1. Title and project info
  2. Documents reviewed (with seniority ranking)
  3. Findings (errors found, with source and resolution)
  4. QA summary table
  5. Professional recommendation
  6. Coordinate table
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT


def generate_survey_report(report_data: dict, output_path: str) -> str:
    doc = Document()

    doc.add_heading(f"Boundary Survey Report", level=0)
    doc.add_heading(f"{report_data['lot_name']} — {report_data['plan_number']}", level=1)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_data = [
        ("Location", report_data.get("location", "")),
        ("Client", report_data.get("client", "")),
        ("Date", report_data.get("date", "")),
        ("Plan Number", report_data.get("plan_number", "")),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
    doc.add_paragraph("")

    doc.add_heading("Documents Reviewed", level=2)
    docs_table = doc.add_table(rows=1, cols=3)
    docs_table.style = "Table Grid"
    hdr = docs_table.rows[0].cells
    hdr[0].text = "Document"
    hdr[1].text = "Type"
    hdr[2].text = "Year"
    for d in report_data.get("documents_reviewed", []):
        row = docs_table.add_row().cells
        row[0].text = d["title"]
        row[1].text = d["type"]
        row[2].text = str(d["year"])

    doc.add_heading("Findings and Errors", level=2)
    errors = report_data.get("errors_found", [])
    if errors:
        errors_table = doc.add_table(rows=1, cols=3)
        errors_table.style = "Table Grid"
        hdr = errors_table.rows[0].cells
        hdr[0].text = "Description"
        hdr[1].text = "Source"
        hdr[2].text = "Resolution"
        for e in errors:
            row = errors_table.add_row().cells
            row[0].text = e["description"]
            row[1].text = e["source"]
            row[2].text = e["resolution"]
    else:
        doc.add_paragraph("No errors or discrepancies found.")

    doc.add_heading("QA Verification Summary", level=2)
    qa_table = doc.add_table(rows=1, cols=3)
    qa_table.style = "Table Grid"
    hdr = qa_table.rows[0].cells
    hdr[0].text = "Check"
    hdr[1].text = "Result"
    hdr[2].text = "Details"
    for q in report_data.get("qa_summary", []):
        row = qa_table.add_row().cells
        row[0].text = q["check"]
        row[1].text = q["result"]
        row[2].text = q["details"]

    doc.add_heading("Professional Recommendation", level=2)
    doc.add_paragraph(report_data.get("recommendation", ""))

    doc.add_heading("Boundary Coordinates", level=2)
    coord_table = doc.add_table(rows=1, cols=4)
    coord_table.style = "Table Grid"
    hdr = coord_table.rows[0].cells
    hdr[0].text = "Corner"
    hdr[1].text = "Easting (m)"
    hdr[2].text = "Northing (m)"
    hdr[3].text = "Monument"
    for c in report_data.get("coordinates", []):
        row = coord_table.add_row().cells
        row[0].text = str(c["corner"])
        row[1].text = f"{c['easting']:.3f}"
        row[2].text = f"{c['northing']:.3f}"
        row[3].text = c.get("monument", "")

    doc.save(output_path)
    return output_path
